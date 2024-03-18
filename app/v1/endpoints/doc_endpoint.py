from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from fastapi.responses import JSONResponse
from pathlib import Path
from docx import Document
from spire.doc import *
from spire.doc.common import *
import os
import uuid
import shutil
from jinja2 import Template


router = APIRouter()
DOCX_WORKSPACE = os.path.abspath("app\\v1\\static\\docx\\")
BASE_URL = "http://192.168.1.44:8000/"


class DocxFilePath(BaseModel):
    docx_file_1_path: str
    docx_file_2_path: str

class RemoveDocxSession(BaseModel):
    session_id: str

class DocxComparator:
    def __init__(self, file_paths: DocxFilePath) -> None:
        self.document_1 = file_paths.docx_file_1_path
        self.document_2 = file_paths.docx_file_2_path

    def validate_documents(self):
        if not os.path.isfile(self.document_1):
            raise HTTPException(status_code=404, detail=f"{self.document_1} not found.")
        if not os.path.isfile(self.document_2):
            raise HTTPException(status_code=404, detail=f"{self.document_2} not found.")
        if not self.validate_docx_format():
            raise HTTPException(status_code=500, detail=f"Invalid docx document format")

    def validate_docx_format(self) -> bool:
        try:
            document_1 = Document(self.document_1)
            document_2 = Document(self.document_2)
            return True
        except Exception as error:
            return False

class Workspace:
    @staticmethod
    def create_session_workspace():
        session_id = str(uuid.uuid4())
        session_folder = os.path.join(DOCX_WORKSPACE, session_id)
        os.makedirs(session_folder)
        return session_id

    @staticmethod
    def copy_documents_to_session_workspace(file1_path: str, file2_path: str, session_workspace_path: str):
        try:
            shutil.copy(file1_path, session_workspace_path)
            shutil.copy(file2_path, session_workspace_path)
            return True
        except Exception as error:
            return False

class HtmlGenerator:
    @staticmethod
    def generate_html_file(session_path, session_id, title, file1, file2):
        template_str = '''<!DOCTYPE html>
        <html>
            <head>
                <title>{{ title }}</title>
                <!-- Bootstrap CSS -->
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.0.2/dist/css/bootstrap.min.css" rel="stylesheet" integrity="sha384-EVSTQN3/azprG1Anm3QDgpJLIm9Nao0Yz1ztcQTwFspd3yD65VohhpuuCOmLASjC" crossorigin="anonymous">
                <script src="https://ajax.googleapis.com/ajax/libs/jquery/3.7.1/jquery.min.js"></script>
                <style>
                .table-responsive {
                    max-height: 600px;
                    overflow-y: auto;
                    overflow-x: auto;
                }
                .table-responsive table {
                    width: 100%;
                }
                .table-container {
                    display: flex;
                    overflow-x: auto;
                }
                .table-container .table-responsive {
                    flex: 0 0 auto;
                    margin-right: 10px;
                }
                .square-badge {
                    border-radius: 0;
                }
                .divScrollDiv {
                    display: inline-block;
                    width: 100%;
                    border: 1px solid black;
                    height: 94vh;
                    overflow: scroll;
                }
                .tableNoScroll {
                    overflow: hidden;
                }
            </style>
        </head>
        <body>
        <div class="col-lg mx-auto p-3 py-md-3">
            <header class="d-flex align-items-center pb-3 mb-5 border-bottom">
                <a href="#" class="d-flex align-items-center text-dark text-decoration-none">
                    <img src="/static/images/logo.jpeg" width="32" height="32" class="p-1">
                    <span class="fs-6">Document Comparison and Analysis (Demo)</span>
                </a>
            </header>
            <div class="container-fluid">
                <div class="row">
                    <div class="col">
                        <div class="table-container">
                            <!-- First Document -->
                            <div class="table table-responsive">
                                <span class="badge bg-primary square-badge mb-3">Docx Document Path</span><span class="badge bg-success square-badge">{{ file1 }}</span>
                                <iframe src="/static/docx/{{session_id}}/original.html" style="width:100%; height:500px;"></iframe>
                            </div> 
                        </div>
                    </div>
                    <div class="col">
                        <div class="table-container">
                            <div class="table table-responsive">
                                <!-- Second Document -->
                                <span class="badge bg-primary square-badge mb-3">Excel Document Path</span><span class="badge bg-success square-badge">{{ file2 }}</span><br>
                                <iframe src="/static/docx/{{session_id}}/result.html" style="width:100%; height:500px;"></iframe>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        <!-- Bootstrap JS (optional) -->
        <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.0.2/dist/js/bootstrap.bundle.min.js" integrity="sha384-MrcW6ZMFYlzcLA8Nl+NtUVF0sA7MsXsP1UyJoMp4YLEuNSfAP+JcXn/tWtIaxVXM" crossorigin="anonymous"></script>
    </body>
    </html>'''
        template = Template(template_str)
        rendered_html = template.render(
            title=title,
            session_path = session_path,
            session_id = session_id,
            file1=file1,
            file2=file2
        )
        with open(f"{session_path}/comparison_result.html", "w") as file:
            file.write(rendered_html)

@router.post("/remove_docx_session")
async def remove_docx_session(sessionid: RemoveDocxSession):
    docx_session_id = sessionid.session_id
    
    """check if session workspace is available"""
    SESSION_WORKSPACE = os.path.join(DOCX_WORKSPACE, docx_session_id)
    if not Path(SESSION_WORKSPACE):
        raise HTTPException(status_code=404, detail=f"Session id {docx_session_id} not available")
    else:
        data = {"message": f"Session id {docx_session_id} removed successfully."}
        try:
            shutil.rmtree(SESSION_WORKSPACE)
            return JSONResponse(content=data, status_code=200)
        except Exception as error:
            raise HTTPException(status_code=500, detail=f"Session id {docx_session_id} not available")

@router.post("/generate_url_for_docx")
async def generate_url(file_paths: DocxFilePath):
    comparator = DocxComparator(file_paths)

    """validate document"""
    comparator.validate_documents()

    """create session workspace"""
    session_id = Workspace.create_session_workspace()
    session_workspace = os.path.join(DOCX_WORKSPACE, session_id)

    if not Workspace.copy_documents_to_session_workspace(file_paths.docx_file_1_path,
                                                         file_paths.docx_file_2_path,
                                                         session_workspace):
        raise HTTPException(status_code=500, detail="Error while copying the documents to session workspace")

    """updated document path"""
    new_doc1_path = os.path.join(session_workspace, os.path.basename(file_paths.docx_file_1_path))
    new_doc2_path = os.path.join(session_workspace, os.path.basename(file_paths.docx_file_2_path))

    """load the first document while initializing the Document object"""
    firstDoc = Document(new_doc1_path)

    """convert the first document to HTML format"""
    original_html_file = rf"{session_workspace}\original.html"
    firstDoc.SaveToFile(original_html_file, FileFormat.Html)

    """load the second document while initializing the Document object"""
    secondDoc = Document(new_doc2_path )

    """compare documents"""
    firstDoc.Compare(secondDoc, "E-ICEBLUE")

    """save comparision result in HTML format"""
    result_file = rf"{session_workspace}\result.html"
    firstDoc.SaveToFile(result_file, FileFormat.Html)

    """generate html"""
    generate_html = HtmlGenerator()
    generate_html.generate_html_file(session_workspace, session_id, "Contentverse Docx Document Comparision", file_paths.docx_file_1_path, 
                                     file_paths.docx_file_2_path)
    
    """generate URL"""
    comparison_result_url = f"{BASE_URL}static/docx/{session_id}/comparison_result.html"
    
    return {"session_id": session_id, "comparison_result_url": comparison_result_url}


